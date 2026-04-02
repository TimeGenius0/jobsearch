#!/usr/bin/env python3
"""
Cover letter generator.

Usage:
    python cover_letter.py <job_url> [--notes "your notes"] [--output cover_letter.docx]
    python cover_letter.py <job_url> --notes "targeting ML infra role, emphasise Attila AI"

Fetches the job description from the URL, then uses Claude to write a crisp,
personalised cover letter based on the embedded resume and any extra notes.
Outputs a formatted .docx file using the template.
"""

import argparse
import asyncio
import os
import re
import sys
from pathlib import Path
from datetime import datetime
from urllib.parse import urlparse

from activity_logger import log_activity, get_recent_entries

import httpx
from docx import Document
from docx.shared import Pt
from langchain_anthropic import ChatAnthropic
from langchain_core.messages import HumanMessage, SystemMessage

try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).resolve().parent / ".env")
except ImportError:
    pass


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = SCRIPT_DIR / "config" / "template" / "Bilel_BOURAOUI_Cover_Letter.docx"
OUTPUT_DIR = SCRIPT_DIR / "output"


# ---------------------------------------------------------------------------
# Resume
# ---------------------------------------------------------------------------
RESUME = """
Product, Science & Engineering Leader

AI founder and product leader with 15+ years building AI products for consumer
services, B2B and enterprise clients. In early-stage startups like Attila AI and
large-scale enterprises like Intuit, I've delivered AI solutions that move the
needle including a marketing campaign recommender that drove 35% omnichannel
penetration growth and a Virtual Sales Assistant that increased cross-sell rates
by 82%. I hold an MSc in Software Management and Electrical Engineering, have led
teams of 40+, and regularly partner with C-suite stakeholders. Also an early-stage
investor, advisor, and aspiring artist.

PROFESSIONAL HIGHLIGHTS

Principal Product Manager, Intuit Mailchimp, Mountain View, CA  (Apr 2025 – Present)
NASDAQ: INTU. 2025 Sales: $18B, employees: 18,000. Hired to report to the CPO.
- Owned Agentic AI strategy: done-for-you campaign recommendations, multimodal
  content generation (layout, text, image), and conversational AI agents.
- Delivered 35% lift in omnichannel penetration growth → +$10M ARR opportunity.
- Led Mailchimp's ChatGPT and Claude apps; built the chat-based agentic AI platform.
- 0→1 agentic content generation to PMF, reducing churn by 78 bps in 1 quarter.
- Established AI quality standards, Evals frameworks, GenAI Trust & Safety guardrails.
- Enabled launch of 6 agentic AI experiences; mentored 8 Principal–Staff PMs on AI
  quality, compliance, and platform integration.

Co-Founder & Chief Product Scientist, AICTOS, San Francisco, CA  (Mar 2023 – May 2025)
AI bootstrapped startup, 2024 Sales: $1.2M, employees: 2.
- Founded and led product for Aictos, a marketplace connecting SMBs with software
  engineering leaders assisted by AI agents for technical leadership and code quality.
- Designed AI agents for technical decision-making, engineering productivity, and
  project risk detection across distributed teams.
- Built a community of 45 experienced CTOs to validate use cases and refine workflows.

Director of Product, ParagonOne / Extern.com (YC 2017), San Francisco, CA  (Apr 2022 – Feb 2023)
Edtech startup backed by YC, 2022 Sales: $8M, employees: 30. Reported to the CTO.
- Led product vision, strategy, and roadmap; overseeing 2 PMs and 5 engineers.
- Pivoted the company to Enterprise/ESG segment after pitching the founders directly.
- Improved SQL conversion rates from 7% to 44% through GTM realignment.

Co-Founder & CPO, Attila AI, Tunis, Tunisia  (Feb 2016 – Jul 2020)
MarketingTech AI startup, 2020 Sales: $3M, employees: 12. Acquired by Craft Foundry.
- Led company from 0→1; hired ~20-person cross-functional team.
- Built ML recommendation engine for banking apps (credit cards, loans, savings).
- Boosted cross-sell conversion rate by 82% vs rule-based systems; +5% in conversions.

VP of Product & Growth, Tayara.tn (subsidiary of Schibsted), Tunis, Tunisia  (Jan 2015 – Jan 2016)
EURONEXT: VEND, 2015 Sales: $1.9B. Hired to report to the head of emerging markets.
- Scaled to 2.5M MAU; 5th most visited site in Tunisia.
- Led team of 10; partnered with 40-person global cross-functional team.

Co-Founder & CPO, Zouz, Tunis, Tunisia  (Mar 2010 – Jun 2014)
Social discovery startup backed by ACP and ATD capital, 2014 Sales: $0.5M, employees: 15.
- Raised seed and Series A funding; grew to 100K MAU.
- 42% activation ratio; 71% 3-month retention.

Product Manager, Sequans Communications, Cupertino, CA  (Apr 2007 – Feb 2010)
NYSE: SQNS — fabless semiconductor (4G/WiMAX SoCs).
- Managed VoIP integration for 4G semiconductor chip.
- Defined GTM strategy for 802.16e SoC; wrote technical white papers on MIMO systems.

EDUCATION
ENST Telecom ParisTech, MS in Electrical Engineering, Paris, France  (2004–2007)

SKILLS
Machine Learning: ML product lifecycle, Transformers/LSTMs, GenAI (GANs, VAEs,
LLMs), MLOps, Edge ML, Neural codecs (EnCodec, SoundStream), data engineering.
"""

SYSTEM_PROMPT = """You are an expert cover letter writer for senior tech and product leaders.
You write crisp, confident, and specific letters — never generic, never sycophantic.
The letter should be 3 short paragraphs:
1. Why this company and role specifically (tie to the JD).
2. Two or three concrete achievements from the resume that directly match the role's needs.
3. A brief, direct close — what you bring and why you'd like to chat.

Tone: direct, warm, no fluff. No "I am writing to express my interest". No bullet points.
Length: under 200 words. Output the letter only — no subject line, no metadata."""

FIT_ASSESSMENT_SYSTEM = """You are evaluating a job description against a candidate's resume.
Write 2–3 sentences assessing fit: what aligns well, what gaps exist, and your overall impression.
Be specific and direct. Output the assessment only — no preamble, no bullet points."""


GRADE_SYSTEM = """You are grading a job opportunity relative to others a candidate has considered.
Grade scale: 5 = excellent fit, strong chance of success; 1 = poor fit, low chance of success.

You will receive the current job's fit assessment and up to 10 previous entries for calibration.
If no previous entries are provided, grade on absolute criteria against the resume.
Respond with JSON only, exactly: {"grade": <integer 1-5>, "rationale": "<one sentence>"}"""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

async def fetch_page_text(url: str) -> str:
    """Fetch a URL and return visible text (HTML stripped via simple heuristic)."""
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/120.0 Safari/537.36"
        )
    }
    async with httpx.AsyncClient(follow_redirects=True, timeout=30) as client:
        resp = await client.get(url, headers=headers)
        resp.raise_for_status()
        content_type = resp.headers.get("content-type", "")
        if "html" in content_type:
            # Strip tags with a simple regex-free approach via html.parser
            from html.parser import HTMLParser

            class _Stripper(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self._chunks = []
                    self._skip = False

                def handle_starttag(self, tag, attrs):
                    if tag in ("script", "style", "nav", "footer", "head"):
                        self._skip = True

                def handle_endtag(self, tag):
                    if tag in ("script", "style", "nav", "footer", "head"):
                        self._skip = False

                def handle_data(self, data):
                    if not self._skip:
                        stripped = data.strip()
                        if stripped:
                            self._chunks.append(stripped)

                def get_text(self):
                    return "\n".join(self._chunks)

            stripper = _Stripper()
            stripper.feed(resp.text)
            return stripper.get_text()[:8000]
        return resp.text[:8000]


def extract_company_from_url(url: str) -> str:
    """
    Extract company name from common job board URL patterns.
    
    Examples:
        https://jobs.lever.co/clickup/role → ClickUp
        https://boards.greenhouse.io/stripe/jobs/123 → Stripe
        https://example.com/careers → Example
    """
    parsed = urlparse(url)
    domain = parsed.netloc.lower()
    path = parsed.path.lower()
    
    # Lever: jobs.lever.co/COMPANY
    if 'lever.co' in domain:
        match = re.search(r'/([^/]+)/', path)
        if match:
            return match.group(1).replace('-', ' ').title()
    
    # Greenhouse: boards.greenhouse.io/COMPANY
    if 'greenhouse.io' in domain:
        match = re.search(r'/([^/]+)/', path)
        if match:
            return match.group(1).replace('-', ' ').title()
    
    # Workday: COMPANY.wd1.myworkdayjobs.com
    if 'myworkdayjobs.com' in domain:
        match = re.search(r'^([^.]+)', domain)
        if match:
            return match.group(1).replace('-', ' ').title()
    
    # Ashby: jobs.ashbyhq.com/COMPANY
    if 'ashbyhq.com' in domain:
        match = re.search(r'/([^/]+)', path)
        if match:
            return match.group(1).replace('-', ' ').title()
    
    # BambooHR: COMPANY.bamboohr.com
    if 'bamboohr.com' in domain:
        match = re.search(r'^([^.]+)', domain)
        if match:
            return match.group(1).replace('-', ' ').title()
    
    # Generic: extract from domain
    # Remove common TLDs and subdomains
    company = domain.replace('www.', '').replace('jobs.', '').replace('careers.', '')
    company = company.split('.')[0]
    return company.replace('-', ' ').title()


def _get_llm() -> ChatAnthropic:
    key = os.getenv("CLAUDE_API_KEY")
    if not key:
        raise ValueError("CLAUDE_API_KEY not set. Add it to .env.")
    return ChatAnthropic(model="claude-sonnet-4-6", temperature=0.7, max_tokens=512, api_key=key)


async def generate_fit_assessment(job_text: str, notes: str) -> str:
    user_content = f"""JOB DESCRIPTION:
{job_text}

MY RESUME:
{RESUME}
"""
    if notes.strip():
        user_content += f"\nEXTRA CONTEXT:\n{notes.strip()}\n"
    user_content += "\nAssess fit now."

    llm = _get_llm()
    response = await llm.ainvoke([
        SystemMessage(content=FIT_ASSESSMENT_SYSTEM),
        HumanMessage(content=user_content),
    ])
    return response.content.strip()


async def grade_job(fit_summary: str, recent_entries: list[dict]) -> tuple[int, str]:
    import json as _json

    context = ""
    if recent_entries:
        lines = []
        for i, entry in enumerate(recent_entries, 1):
            grade_str = f" — grade: {entry['grade']}" if entry.get("grade") else ""
            lines.append(
                f"{i}. {entry.get('company', 'Unknown')} — \"{entry['fit_summary']}\"{grade_str}"
            )
        context = "\n\nPrevious jobs for calibration:\n" + "\n".join(lines)

    user_content = f"Current job fit assessment:\n{fit_summary}{context}\n\nGrade the current job now."

    llm = _get_llm()
    response = await llm.ainvoke([
        SystemMessage(content=GRADE_SYSTEM),
        HumanMessage(content=user_content),
    ])
    parsed = _json.loads(response.content.strip())
    return int(parsed["grade"]), str(parsed["rationale"])


async def generate_cover_letter(job_url: str, job_text: str, notes: str) -> str:
    user_content = f"""JOB DESCRIPTION (fetched from {job_url}):
{job_text}

MY RESUME:
{RESUME}
"""
    if notes.strip():
        user_content += f"\nMY NOTES / EXTRA CONTEXT:\n{notes.strip()}\n"

    user_content += "\nWrite the cover letter now."

    llm = _get_llm()
    response = await llm.ainvoke([
        SystemMessage(content=SYSTEM_PROMPT),
        HumanMessage(content=user_content),
    ])
    return response.content.strip()


def create_formatted_docx(letter_text: str, output_path: Path, company_name: str = None) -> None:
    """
    Create a formatted .docx using the template.
    Replaces the letter body (paragraphs 4-11) with generated content.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
    
    # Load template
    doc = Document(str(TEMPLATE_PATH))
    
    # Keep greeting generic (already "Dear Hiring Manager," in template)
    # No customization needed
    
    # Delete old letter body (paragraphs 6, 8, 10 - the content paragraphs)
    # Keep blank paragraphs for spacing (5, 7, 9, 11)
    # Work backwards to avoid index shifting
    paragraphs_to_replace = [10, 8, 6]  # Content paragraphs in reverse order
    
    for idx in paragraphs_to_replace:
        if len(doc.paragraphs) > idx:
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)
    
    # Now insert new content paragraphs
    # Split letter into paragraphs
    letter_paragraphs = [p.strip() for p in letter_text.split("\n\n") if p.strip()]
    
    # Insert after greeting (after paragraph 4)
    # We'll insert: blank, para1, blank, para2, blank, para3, blank
    insert_position = 5  # After greeting
    
    for i, para_text in enumerate(letter_paragraphs):
        # Add blank line
        blank = doc.add_paragraph()
        doc._body._element.insert(insert_position, blank._element)
        insert_position += 1
        
        # Add content paragraph with Cambria 11pt formatting
        new_para = doc.add_paragraph()
        run = new_para.add_run(para_text)
        run.font.name = 'Cambria'
        run.font.size = Pt(11)
        doc._body._element.insert(insert_position, new_para._element)
        insert_position += 1
    
    # Add final blank line before signature
    blank = doc.add_paragraph()
    doc._body._element.insert(insert_position, blank._element)
    
    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"\n✅ Cover letter saved to: {output_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

async def _run_pipeline(
    job_url: str, notes: str, company_name: str, output_path: Path
) -> tuple[str, str, int | None]:
    """Fetch JD, run all 3 Claude calls, print results, return (letter, fit_summary, grade)."""
    job_text = await fetch_page_text(job_url)

    # Call 1 — cover letter
    letter = await generate_cover_letter(job_url, job_text, notes)
    print("=" * 60)
    print(letter)
    print("=" * 60)
    create_formatted_docx(letter, output_path, company_name)

    # Call 2 — fit assessment
    fit_summary = ""
    try:
        fit_summary = await generate_fit_assessment(job_text, notes)
        print(f"\n🎯 Fit Assessment:\n{fit_summary}\n")
    except Exception as e:
        print(f"\n⚠️  Fit assessment failed: {e}")
        return letter, fit_summary, None

    # Call 3 — grade
    grade = None
    try:
        recent = get_recent_entries(10)
        grade, rationale = await grade_job(fit_summary, recent)
        print(f"📊 Grade: {grade}/5\n{rationale}\n")
    except Exception as e:
        print(f"\n⚠️  Grading failed: {e}")

    return letter, fit_summary, grade


def main():
    parser = argparse.ArgumentParser(description="Generate a formatted cover letter for a job posting.")
    parser.add_argument("url", help="URL of the job description")
    parser.add_argument("--notes", default="", help="Extra context or targeting notes")
    parser.add_argument("--output", help="Output .docx filename (default: auto-generated)")
    parser.add_argument("--company", help="Company name (optional, override auto-detection)")
    args = parser.parse_args()

    # Clean URL (remove whitespace/newlines from copy-paste)
    job_url = args.url.strip()

    # Extract company name from URL if not provided
    company_name = args.company if args.company else extract_company_from_url(job_url)
    print(f"📋 Company: {company_name}")
    print(f"🔗 Fetching job description from {job_url} …\n")
    
    # Generate output filename
    if args.output:
        output_path = OUTPUT_DIR / args.output
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        company_slug = company_name.replace(" ", "_")
        output_path = OUTPUT_DIR / f"Cover_Letter_{company_slug}_{timestamp}.docx"

    try:
        _, fit_summary, grade = asyncio.run(
            _run_pipeline(job_url, args.notes, company_name, output_path)
        )
        log_activity(
            "cover-letter", job_url, company_name, "success", str(output_path),
            fit_summary=fit_summary, grade=grade,
        )
    except Exception as e:
        log_activity("cover-letter", job_url, company_name, "error", str(e))
        raise


if __name__ == "__main__":
    main()
